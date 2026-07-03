"""DOCX parser (optional dependency: python-docx).

Install with:  pip install "agentcontext[docx]"
"""

from __future__ import annotations

from ..core.model import Block, BlockType, Document, Provenance
from .base import Parser, register_parser


class DocxParser(Parser):
    name = "docx"
    extensions = ("docx",)

    def parse(self, path: str) -> Document:
        try:
            import docx  # python-docx
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ImportError(
                "DOCX parsing needs python-docx. Install with: pip install \"agentcontext[docx]\""
            ) from exc

        d = docx.Document(path)
        doc = Document(source=path, meta={"parser": self.name})
        section = ""
        for para in d.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            prov = Provenance(source=path, section=section, parser=self.name, version="docx-parser/0.1")
            if style.startswith("heading"):
                level = int("".join(filter(str.isdigit, style)) or 1)
                if level <= 1:
                    section = text
                doc.add(Block(type=BlockType.HEADING, text=text, level=level, provenance=prov))
            elif style.startswith("list") or para.text.startswith(("- ", "•")):
                doc.add(Block(type=BlockType.LIST_ITEM, text=text.lstrip("-• "), provenance=prov))
            else:
                doc.add(Block(type=BlockType.PARAGRAPH, text=text, provenance=prov))

        # Tables -> markdown-rendered TABLE blocks (kept in reading order at end for v0.1)
        for table in d.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            md = _rows_to_markdown(rows)
            doc.add(
                Block(
                    type=BlockType.TABLE,
                    text=md,
                    meta={"markdown": md, "rows": rows},
                    provenance=Provenance(source=path, section=section, parser=self.name),
                )
            )
        return doc


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header, *body = rows
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join("---" for _ in header) + " |"]
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


register_parser(DocxParser())
